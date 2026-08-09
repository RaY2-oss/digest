# -*- coding: utf-8 -*-
"""
word_generator.py — сборка готового .docx из списка пересказанных статей.

Единственная публичная функция — build_digest(summaries), вызывается из
sunday_processor_mmr.py после того, как все статьи пересказаны и прошли
валидацию. Результат сохраняется в config.OUTPUT_DIR/digest_<дата>.docx.
"""
import os
import time
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import config


def _parse_date(raw: str) -> str:
    """
    Конвертирует ISO-дату из БД (YYYY-MM-DD или YYYY-MM-DD HH:MM:SS)
    в формат дд.мм.гггг. При неудаче возвращает исходную строку.
    """
    if not raw:
        return ""
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(raw.strip()[:len(fmt)], fmt).strftime("%d.%m.%Y")
        except ValueError:
            continue
    try:
        return datetime.strptime(raw.strip()[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        return raw.strip()


def _add_horizontal_rule(paragraph):
    """python-docx не умеет рисовать горизонтальную линию напрямую —
    добавляем нижнюю границу параграфу через сырой OOXML (w:pBdr/w:bottom)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_digest(summaries):
    """
    summaries — список словарей {"title","summary","url","publish_date"}
    и необязательный "urls" — все источники сюжета, по которым сделан
    пересказ (перепечатки одной новости в разных изданиях).
    publish_date — строка "YYYY-MM-DD" или None.
    """
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    doc = Document()

    doc.add_heading("Еженедельный новостной дайджест", level=0)

    sub = doc.add_paragraph()
    run = sub.add_run(
        "Тематика: наука, образование и молодёжная политика. "
        "Регион: Турция, Центральная Азия, Южный Кавказ"
    )
    run.italic = True
    run.font.size = Pt(11)

    date_str = time.strftime("%d.%m.%Y")
    doc.add_paragraph(f"Дата формирования: {date_str}")
    _add_horizontal_rule(doc.add_paragraph())

    for i, item in enumerate(summaries, start=1):
        title        = item.get("title", "Без заголовка")
        summary      = item.get("summary", "")
        urls         = item.get("urls") or [item.get("url", "")]
        publish_date = item.get("publish_date") or ""

        # Заголовок: "N. дд.мм.гггг — Текст заголовка"
        p_title = doc.add_paragraph()
        r_num = p_title.add_run(f"{i}. ")
        r_num.bold = True
        r_num.font.size = Pt(12)

        formatted_date = _parse_date(publish_date) if publish_date else ""
        if formatted_date:
            r_date = p_title.add_run(f"{formatted_date} — ")
            r_date.bold = False
            r_date.italic = True
            r_date.font.size = Pt(10)

        r_title = p_title.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(12)

        p_sum = doc.add_paragraph()
        r_sum = p_sum.add_run(summary)
        r_sum.font.size = Pt(11)

        # Источников может быть несколько: пересказ сюжета собирается по всем
        # перепечаткам, попавшим в промпт, и все они должны быть проверяемы.
        # Идут одной строкой через «; »: с новой строки каждый занимал под
        # сюжетом больше места, чем сам пересказ.
        live = [u for u in urls if u]
        if live:
            p_url = doc.add_paragraph()
            r_url = p_url.add_run("URL: " + "; ".join(live))
            r_url.italic = True
            r_url.font.size = Pt(10)

        doc.add_paragraph()

    out_name = f"digest_{time.strftime('%Y-%m-%d')}.docx"
    out_path = os.path.join(config.OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_path
