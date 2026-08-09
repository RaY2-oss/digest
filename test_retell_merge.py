# -*- coding: utf-8 -*-
"""
test_retell_merge.py — пересказ сюжета по НЕСКОЛЬКИМ перепечаткам сразу.

Раньше в промпт уходила ровно одна статья, а её дубликаты с других сайтов
_process_slot просто выбрасывал (прогон 26.07: 49 отброшенных). Детали, которые
сохранила только одна редакция — цифры, имена, цитаты, — терялись вместе с ними.
Потом версии стали склеиваться в один промпт, но каждая резалась по первым N
символов — а перепечатки начинаются одинаково, и модель получала один и тот же
лид четыре раза. Теперь в промпт идёт ОТБОР по всем версиям сразу
(retell_select): предложения, которых ещё не было, от как можно большего числа
изданий. Все использованные ссылки идут в документ.

Запуск: venv/bin/python test_retell_merge.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import config
import sunday_processor_mmr as spm
import word_generator

_STEM = "alpha beta gamma delta epsilon zeta eta theta iota kappa".split()


def _sent(seed):
    """Предложение из слов, которых нет больше нигде: отбор обязан засчитать
    его как новый факт, а не как пересказ соседнего."""
    return " ".join("%s%d" % (w, seed) for w in _STEM) + "."


def _text(seed, n_sents):
    return " ".join(_sent(seed * 1000 + i) for i in range(n_sents))


def test_single_version_fills_the_whole_budget():
    """Одна версия получает весь бюджет промпта, а не свою долю от него."""
    src, _ = spm._retell_sources([("https://a.com/1", _text(1, 200))])
    assert "https://a.com/1" in src
    body = src.split("Article text:\n", 1)[1]
    assert config.RETELL_CHARS_TOTAL * 0.8 < len(body) <= config.RETELL_CHARS_TOTAL


def test_several_versions_share_the_budget():
    versions = [(f"https://o{i}.com/1", _text(i, 200)) for i in range(4)]
    src, _ = spm._retell_sources(versions)
    assert len(src) <= config.RETELL_CHARS_TOTAL * 1.3, len(src)   # + заголовки блоков
    for i in range(4):
        assert f"https://o{i}.com/1" in src
    assert src.count("----------") == 3, "версии обязаны быть разделены"


def test_reprints_do_not_eat_the_budget_twice():
    """Главное, ради чего отбор и делался: пять дословных перепечаток обязаны
    занять в промпте столько же места, сколько одна статья, — иначе модель
    четырежды читает один и тот же лид вместо четырёх разных деталей."""
    one, _ = spm._retell_sources([("https://o0.com/1", _text(0, 12))])
    five, _ = spm._retell_sources([(f"https://o{i}.com/1", _text(0, 12)) for i in range(5)])
    body_one = one.split("Article text:\n", 1)[1]
    bodies = five.split("Article text:\n")[1:]
    assert sum(len(b) for b in bodies) < len(body_one) * 1.5, [len(b) for b in bodies]


def test_extra_versions_are_dropped_not_squeezed():
    """Потолок по числу версий важнее полноты: разбирать двадцать перепечаток
    ради одного-двух новых предложений не стоит времени."""
    versions = [(f"https://o{i}.com/1", _text(i, 40)) for i in range(10)]
    src, _ = spm._retell_sources(versions)
    assert f"https://o{config.RETELL_MAX_VERSIONS - 1}.com/1" in src
    assert f"https://o{config.RETELL_MAX_VERSIONS}.com/1" not in src


def test_empty_text_does_not_break_the_prompt():
    src, urls = spm._retell_sources([("https://a.com/1", None),
                                     ("https://b.com/1", _text(2, 5))])
    assert "https://a.com/1" not in src, "версии без текста в промпте нечего делать"
    assert "https://b.com/1" in src
    assert urls == ["https://b.com/1"], "источником зовётся только прочитанный"


def test_nothing_to_pick_gives_an_empty_prompt():
    """Все версии — мусор без единого предложения: промпт пустой, а не битый."""
    assert spm._retell_sources([("https://a.com/1", "5 30 2026"),
                                ("https://b.com/1", "")]) == ("", [])


def test_docx_lists_every_source_url(tmp="/tmp"):
    """Все ссылки, попавшие в промпт, обязаны быть в документе — иначе пересказ
    нечем проверить."""
    old_out = config.OUTPUT_DIR
    config.OUTPUT_DIR = tmp
    try:
        path = word_generator.build_digest([{
            "title": "Заголовок", "summary": "Пересказ.",
            "url": "https://a.com/1",
            "urls": ["https://a.com/1", "https://b.com/1", "https://c.com/1"],
            "publish_date": "2026-07-27",
        }])
        from docx import Document
        paras = [p.text for p in Document(path).paragraphs]
        text = "\n".join(paras)
        for u in ("https://a.com/1", "https://b.com/1", "https://c.com/1"):
            assert u in text, f"{u} потерян в документе"
        # ...и все в ОДНОЙ строке через точку с запятой: тремя абзацами подряд
        # ссылки занимали под сюжетом больше места, чем сам пересказ.
        line = next(p for p in paras if p.startswith("URL: "))
        assert line == ("URL: https://a.com/1; https://b.com/1; "
                        "https://c.com/1"), line
    finally:
        config.OUTPUT_DIR = old_out


def test_docx_falls_back_to_single_url():
    """Старый формат словаря (без "urls") обязан рисоваться как раньше."""
    old_out = config.OUTPUT_DIR
    config.OUTPUT_DIR = "/tmp"
    try:
        path = word_generator.build_digest([{
            "title": "Заголовок", "summary": "Пересказ.",
            "url": "https://a.com/1", "publish_date": "2026-07-27",
        }])
        from docx import Document
        text = "\n".join(p.text for p in Document(path).paragraphs)
        assert "URL: https://a.com/1" in text
    finally:
        config.OUTPUT_DIR = old_out


def test_story_date_is_the_earliest_reprint():
    """Дату сюжету даёт первая его перепечатка, а не представитель отбора.

    Представителя выбирает MMR, и это запросто третья по счёту публикация:
    датировать событие ею — значит показать в дайджесте день, когда сюжет
    перепечатали, вместо дня, когда о нём написали.
    """
    rep = ("u", "t", None, "2026-08-06", "TR", 0.0, "TR1")
    assert spm._story_date(rep, "нет") == "2026-08-06"

    story = rep + ([("a", "", None, "2026-08-08", "TR", 0.0, "TR1"),
                    ("b", "", None, "2026-08-04", "TR", 0.0, "TR1"),
                    # у части статей publish_date не определился — такие в
                    # счёт не идут, иначе пустая строка выиграла бы min()
                    ("c", "", None, "", "TR", 0.0, "TR1")],)
    assert spm._story_date(story, "нет") == "2026-08-04"

    blank = ("u", "t", None, "", "TR", 0.0, "TR1",
             [("a", "", None, "", "TR", 0.0, "TR1")])
    assert spm._story_date(blank, "2026-08-01") == "2026-08-01"


def test_event_date_wins_only_when_it_checks_out():
    """Дату события называет модель, а печатается она, только если сторож её
    подтвердил: в недельном окне, не позже публикации и буквально стоящая в
    тексте. Иначе в шапке пункта остаётся дата публикации.

    Проверяется весь путь до словаря пункта: событие, о котором написали через
    три дня, и есть тот случай, ради которого поле заведено."""
    import datetime

    import event_date

    day = datetime.date.today()
    ev_d = day - datetime.timedelta(days=3)
    ev = ev_d.isoformat()
    pub = (day - datetime.timedelta(days=1)).isoformat()
    # В тексте дата стоит так, как её пишет издание, — днём вперёд и без
    # ведущего нуля; сторож обязан узнать и такое написание.
    text = ("Konferans %d.%d.%d tarihinde Ankarada duzenlendi ve katilimcilar "
            "ilgi gosterdi bilim insanlari." % (ev_d.day, ev_d.month, ev_d.year))

    real = spm._call_openrouter_raw
    try:
        spm._call_openrouter_raw = lambda s, u, ref_url=None: (
            '{"title": "В Турции прошла конференция по науке", '
            '"summary": "В Анкаре прошла научная конференция исследователей.", '
            '"event_date": "%s"}' % ev)
        got = spm._retell_article("https://a.com/1", [("https://a.com/1", text)],
                                  pub, "тест")
        assert got and got["publish_date"] == ev, got

        # Та же дата, но в тексте её нет — печатается публикация.
        spm._call_openrouter_raw = lambda s, u, ref_url=None: (
            '{"title": "В Турции прошла конференция по науке", '
            '"summary": "В Анкаре прошла научная конференция исследователей.", '
            '"event_date": "%s"}'
            % (day - datetime.timedelta(days=2)).isoformat())
        got = spm._retell_article("https://a.com/1", [("https://a.com/1", text)],
                                  pub, "тест")
        assert got and got["publish_date"] == pub, got
    finally:
        spm._call_openrouter_raw = real

    # Потолок недели держится и без всякой модели.
    assert event_date.clamp((day - datetime.timedelta(days=30)).isoformat(),
                            day) == ""


if __name__ == "__main__":
    test_story_date_is_the_earliest_reprint()
    test_event_date_wins_only_when_it_checks_out()
    test_single_version_fills_the_whole_budget()
    test_several_versions_share_the_budget()
    test_reprints_do_not_eat_the_budget_twice()
    test_extra_versions_are_dropped_not_squeezed()
    test_empty_text_does_not_break_the_prompt()
    test_nothing_to_pick_gives_an_empty_prompt()
    test_docx_lists_every_source_url()
    test_docx_falls_back_to_single_url()
    print("ok")
